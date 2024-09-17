# controllers/report_controller.py
import io
from odoo import http
from odoo.http import request
from docx import Document  # Import docx library to generate DOCX files

class ReportController(http.Controller):

    @http.route('/report/docx/product_profit/<int:report_id>', type='http', auth="user")
    def product_profit_report_docx(self, report_id):
        # Fetch the report data
        report = request.env['product.profit.report'].sudo().browse(report_id)
        profit_data = report.calculate_profit()

        # Generate DOCX
        doc = Document()
        doc.add_heading(f'Product Profit Report - {report.fiscal_year_id.name}', 0)

        # Add table
        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Product Name'
        hdr_cells[1].text = 'Total Sales'
        hdr_cells[2].text = 'Total Purchases'
        hdr_cells[3].text = 'Profit'

        for line in profit_data:
            row_cells = table.add_row().cells
            row_cells[0].text = line['product_name']
            row_cells[1].text = str(line['total_sales'])
            row_cells[2].text = str(line['total_purchases'])
            row_cells[3].text = str(line['profit'])

        # Save DOCX to memory stream
        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)

        # Serve DOCX as a file
        return request.make_response(stream.getvalue(), headers=[
            ('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'),
            ('Content-Disposition', f'attachment; filename=Product_Profit_Report_{report.fiscal_year_id.name}.docx;')
        ])
